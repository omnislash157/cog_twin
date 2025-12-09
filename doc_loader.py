"""
Document Loader - DOCX loading and context stuffing for enterprise mode.

Loads .docx files from hierarchical folder structure, caches content,
and provides division-aware context building.

Folder structure:
    manuals/
    ├── Driscoll/
    │   ├── Warehouse/        # Division folder
    │   │   ├── Dispatching Manual.docx
    │   │   ├── Driver Check-in Manual.docx
    │   │   └── ...
    │   ├── HR/
    │   ├── Purchasing/
    │   └── Shared/           # Accessible by all divisions

Usage:
    from doc_loader import DocLoader, DivisionContextBuilder

    loader = DocLoader(Path("./manuals"))
    stats = loader.get_stats()

    builder = DivisionContextBuilder(loader)
    context = builder.get_context_for_division("warehouse", max_tokens=200000)

Version: 1.0.0
"""

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class LoadedDoc:
    """A loaded document with metadata."""
    path: Path
    name: str
    division: str
    content: str
    char_count: int
    approx_tokens: int
    paragraphs: int

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "division": self.division,
            "char_count": self.char_count,
            "approx_tokens": self.approx_tokens,
            "paragraphs": self.paragraphs,
        }


@dataclass
class DocStats:
    """Statistics about loaded documents."""
    total_docs: int
    total_chars: int
    total_tokens: int
    by_division: Dict[str, Dict[str, int]]
    doc_list: List[str]


# =============================================================================
# DOCX LOADER
# =============================================================================

class DocLoader:
    """
    Loads and caches .docx files from a directory tree.

    Extracts text content and organizes by division (folder structure).
    """

    # Approximate tokens per character (conservative estimate)
    CHARS_PER_TOKEN = 4

    def __init__(self, docs_dir: Path):
        """
        Initialize document loader.

        Args:
            docs_dir: Root directory containing documents
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx required for document loading. "
                "Install with: pip install python-docx"
            )

        self.docs_dir = Path(docs_dir)
        self._cache: Dict[str, LoadedDoc] = {}
        self._loaded = False

    def _extract_text(self, docx_path: Path) -> str:
        """Extract all text from a .docx file."""
        try:
            doc = Document(docx_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)

        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {e}")
            return ""

    def _detect_division(self, docx_path: Path) -> str:
        """
        Detect division from folder path.

        Folder structure relative to docs_dir:
            Division/foo.docx -> "division"

        Examples (docs_dir = manuals/Driscoll):
            Warehouse/foo.docx -> "warehouse"
            HR/bar.docx -> "hr"
            Shared/baz.docx -> "shared"
        """
        # Get path relative to docs_dir
        try:
            rel_path = docx_path.relative_to(self.docs_dir)
            parts = rel_path.parts
        except ValueError:
            return "general"

        # parts[0] = division folder, parts[-1] = filename
        if len(parts) >= 2:
            # File is in a subfolder - that's the division
            division = parts[0].lower()
            return division
        elif len(parts) == 1:
            # File directly in docs_dir (no division folder)
            return "general"

        return "general"

    def _load_all(self):
        """Load all .docx files from docs directory."""
        if self._loaded:
            return

        if not self.docs_dir.exists():
            logger.warning(f"Docs directory not found: {self.docs_dir}")
            self._loaded = True
            return

        # Find all .docx files recursively
        docx_files = list(self.docs_dir.rglob("*.docx"))

        # Filter out temp files (start with ~)
        docx_files = [f for f in docx_files if not f.name.startswith("~")]

        logger.info(f"Found {len(docx_files)} .docx files in {self.docs_dir}")

        for docx_path in docx_files:
            try:
                content = self._extract_text(docx_path)
                if not content:
                    continue

                division = self._detect_division(docx_path)
                char_count = len(content)
                approx_tokens = char_count // self.CHARS_PER_TOKEN
                para_count = content.count("\n\n") + 1

                doc = LoadedDoc(
                    path=docx_path,
                    name=docx_path.stem,  # Filename without extension
                    division=division,
                    content=content,
                    char_count=char_count,
                    approx_tokens=approx_tokens,
                    paragraphs=para_count,
                )

                # Use relative path as key
                key = str(docx_path.relative_to(self.docs_dir))
                self._cache[key] = doc

                logger.debug(f"Loaded: {doc.name} ({division}) - ~{approx_tokens} tokens")

            except Exception as e:
                logger.error(f"Failed to load {docx_path}: {e}")

        self._loaded = True
        logger.info(f"Loaded {len(self._cache)} documents into cache")

    def get_docs_for_division(self, division: str) -> List[LoadedDoc]:
        """
        Get all documents for a division.

        Args:
            division: Division name (e.g., "warehouse", "hr")

        Returns:
            List of LoadedDoc for that division
        """
        self._load_all()

        division_lower = division.lower()
        return [
            doc for doc in self._cache.values()
            if doc.division == division_lower
        ]

    def get_all_docs(self) -> List[LoadedDoc]:
        """Get all loaded documents."""
        self._load_all()
        return list(self._cache.values())

    def get_stats(self) -> DocStats:
        """Get statistics about loaded documents."""
        self._load_all()

        by_division: Dict[str, Dict[str, int]] = {}
        total_chars = 0
        total_tokens = 0
        doc_list = []

        for doc in self._cache.values():
            if doc.division not in by_division:
                by_division[doc.division] = {
                    "docs": 0,
                    "chars": 0,
                    "tokens": 0,
                }

            by_division[doc.division]["docs"] += 1
            by_division[doc.division]["chars"] += doc.char_count
            by_division[doc.division]["tokens"] += doc.approx_tokens

            total_chars += doc.char_count
            total_tokens += doc.approx_tokens
            doc_list.append(f"{doc.division}/{doc.name}")

        return DocStats(
            total_docs=len(self._cache),
            total_chars=total_chars,
            total_tokens=total_tokens,
            by_division=by_division,
            doc_list=doc_list,
        )


# =============================================================================
# CONTEXT BUILDER
# =============================================================================

class DivisionContextBuilder:
    """
    Builds context strings for stuffing into LLM prompts.

    Division-aware: only includes docs relevant to user's division.
    Respects token limits by truncating or selecting subset.
    """

    def __init__(self, docs_dir_or_loader):
        """
        Initialize context builder.

        Args:
            docs_dir_or_loader: Either a Path to docs dir or a DocLoader instance
        """
        if isinstance(docs_dir_or_loader, DocLoader):
            self.loader = docs_dir_or_loader
        else:
            self.loader = DocLoader(Path(docs_dir_or_loader))

    def get_context_for_division(
        self,
        division: str,
        max_tokens: int = 200000,
        include_shared: bool = True,
    ) -> str:
        """
        Get formatted context string for a division.

        Args:
            division: User's division (e.g., "warehouse")
            max_tokens: Maximum tokens to include
            include_shared: Whether to include shared documents

        Returns:
            Formatted context string for prompt injection
        """
        # Get division docs
        docs = self.loader.get_docs_for_division(division)

        # Add shared docs if requested
        if include_shared and division != "shared":
            shared_docs = self.loader.get_docs_for_division("shared")
            docs.extend(shared_docs)

        if not docs:
            logger.warning(f"No documents found for division: {division}")
            return ""

        # Sort by token count (smaller docs first for better fit)
        docs.sort(key=lambda d: d.approx_tokens)

        # Build context respecting token limit
        sections = []
        tokens_used = 0
        docs_included = 0

        # Header
        header = f"=== COMPANY DOCUMENTATION ({division.upper()}) ===\n"
        header += "The following documents are your authoritative source for procedures and policies.\n"
        header += "Cite document names when answering questions.\n\n"

        header_tokens = len(header) // 4
        tokens_used += header_tokens
        sections.append(header)

        for doc in docs:
            # Check if we have room for this doc
            if tokens_used + doc.approx_tokens > max_tokens:
                # Try to fit a truncated version
                remaining_tokens = max_tokens - tokens_used
                if remaining_tokens > 500:  # Worth including a truncated version
                    max_chars = remaining_tokens * 4
                    truncated = doc.content[:max_chars]
                    truncated += "\n\n[DOCUMENT TRUNCATED - ASK FOR SPECIFIC SECTIONS]"

                    section = f"--- {doc.name} (from {doc.division}) ---\n\n"
                    section += truncated
                    section += "\n\n"

                    sections.append(section)
                    docs_included += 1

                break

            # Add full document
            section = f"--- {doc.name} (from {doc.division}) ---\n\n"
            section += doc.content
            section += "\n\n"

            sections.append(section)
            tokens_used += doc.approx_tokens
            docs_included += 1

        # Footer
        footer = f"=== END DOCUMENTATION ({docs_included} documents, ~{tokens_used} tokens) ===\n"
        sections.append(footer)

        context = "".join(sections)
        logger.info(f"Built context for {division}: {docs_included} docs, ~{tokens_used} tokens")

        return context

    def get_context_for_divisions(
        self,
        divisions: List[str],
        max_tokens: int = 200000,
    ) -> str:
        """
        Get context for multiple divisions.

        Useful for managers who need access to multiple areas.
        """
        all_docs = []

        for division in divisions:
            docs = self.loader.get_docs_for_division(division)
            all_docs.extend(docs)

        # Deduplicate (in case "shared" was included multiple times)
        seen_paths = set()
        unique_docs = []
        for doc in all_docs:
            if str(doc.path) not in seen_paths:
                seen_paths.add(str(doc.path))
                unique_docs.append(doc)

        # Sort and build
        unique_docs.sort(key=lambda d: d.approx_tokens)

        sections = []
        tokens_used = 0
        docs_included = 0

        header = f"=== COMPANY DOCUMENTATION (MULTI-DIVISION ACCESS) ===\n"
        header += f"Divisions: {', '.join(divisions)}\n"
        header += "Cite document names when answering questions.\n\n"

        header_tokens = len(header) // 4
        tokens_used += header_tokens
        sections.append(header)

        for doc in unique_docs:
            if tokens_used + doc.approx_tokens > max_tokens:
                break

            section = f"--- {doc.name} [{doc.division}] ---\n\n"
            section += doc.content
            section += "\n\n"

            sections.append(section)
            tokens_used += doc.approx_tokens
            docs_included += 1

        footer = f"=== END DOCUMENTATION ({docs_included} documents) ===\n"
        sections.append(footer)

        return "".join(sections)


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    import sys

    if not DOCX_AVAILABLE:
        print("Error: python-docx required. Run: pip install python-docx")
        sys.exit(1)

    # Default to ./manuals
    docs_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("./manuals")

    if not docs_dir.exists():
        print(f"Directory not found: {docs_dir}")
        sys.exit(1)

    print(f"Loading documents from: {docs_dir}")
    print("=" * 60)

    loader = DocLoader(docs_dir)
    stats = loader.get_stats()

    print(f"\nTotal documents: {stats.total_docs}")
    print(f"Total characters: {stats.total_chars:,}")
    print(f"Approximate tokens: {stats.total_tokens:,}")

    print("\nBy division:")
    for division, div_stats in stats.by_division.items():
        print(f"  {division}:")
        print(f"    Docs: {div_stats['docs']}")
        print(f"    Tokens: {div_stats['tokens']:,}")

    print("\nDocument list:")
    for doc_name in stats.doc_list:
        print(f"  - {doc_name}")

    # Test context building
    if "--context" in sys.argv:
        builder = DivisionContextBuilder(loader)

        # Try warehouse division
        if "warehouse" in stats.by_division:
            print("\n" + "=" * 60)
            print("SAMPLE CONTEXT (warehouse, 50K tokens max)")
            print("=" * 60)
            context = builder.get_context_for_division("warehouse", max_tokens=50000)
            print(context[:2000] + "\n...[truncated]...")
